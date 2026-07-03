import json
import os
import streamlit as st
import pdfplumber
from dotenv import load_dotenv
from langchain_classic.agents import AgentExecutor, create_tool_calling_agent
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from docx import Document as DocxDocument
from langchain.tools import tool

# ============================================================
# 加载环境变量
# ============================================================
load_dotenv()


def load_qa_db():
    """读取题库 JSON 文件，返回 Document 列表"""
    with open("questions.JSON", "r", encoding="utf-8") as f:
        data = json.load(f)

    docs = []

    for item in data:
        content = f"问题:{item['question']}\n 答案:{item['answer']}\n 分类:{item['topic']}"
        meta = {
            "id": item["id"],
            "topic": item["topic"],
            "language": item["language"],
            "difficulty": item["difficulty"],
        }
        doc = Document(page_content=content, metadata=meta)
        docs.append(doc)

    return docs


def init_vectorstore():
    """
    加载题库文档并存入 Chroma，返回 vectorstore 实例。
    使用智谱 AI Embedding API（云端，无需下载模型）。
    """
    # 智谱 AI Embedding，API 兼容 OpenAI 格式
    # model: embedding-2（智谱 embedding 模型）
    # chunk_size=32: 智谱限制每次最多 64 条，分批发送避免 400 错误
    embeddings = OpenAIEmbeddings(
        model="embedding-2",
        api_key=os.getenv("ZHIPUAI_KEY"),
        base_url=os.getenv("ZHIPUAI_BASE_URL"),
        chunk_size=32,
    )

    # 如果已经建库了，直接加载
    if os.path.exists("./Chroma_db") and os.listdir("./Chroma_db"):
        return Chroma(
            persist_directory="./Chroma_db",
            embedding_function=embeddings,
        )

    # 第一次加载才插入
    docs = load_qa_db()
    return Chroma.from_documents(
        documents=docs,
        embedding=embeddings,
        persist_directory="./Chroma_db",
    )

#简历解析函数
def parse_resume(file) -> str :
    """
    解析上传的简历文件(支持PDF/Word/TXT).返回纯文本
    """
    if file.type == "application/pdf":
        with pdfplumber.open(file) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif file.type == "application/vnd.openxmlformats.wordprocessingml.document":
        doc = DocxDocument(file)
        return "\n".join(para.text  for para in doc.paragraphs)
    else:
        #TXT或其他文本文件
        return file.read().decode("utf-8")



#构建检索链
def get_llm(temperature=0.7):
    """主模型:出题,agent推理"""
    return ChatOpenAI(
        model="glm-4-flash",
        api_key=os.getenv("ZHIPUAI_KEY"),
        base_url=os.getenv("ZHIPUAI_BASE_URL"),
        temperature=temperature
    )
def get_eval_llm():
    """评估专用模型:保证评分稳定准确"""
    return ChatOpenAI(
        model="glm-4-flash",
        api_key=os.getenv("ZHIPUAI_KEY"),
        base_url=os.getenv("ZHIPUAI_BASE_URL"),
        temperature=0.1,
    )

#搜题工具
@tool
def search_question_bank(query: str, difficult: str = "medium") -> str:
    """
    从面试题库中检索与 query 相关的问题。
    返回结果包含题目和参考答案——但参考答案仅供 Agent 内部评估使用，
    绝对不要把答案透露给候选人！
    """
    vectorstore = init_vectorstore()
    filter_dict = {"difficulty": difficult} if difficult else {}
    retriever = vectorstore.as_retriever(
        search_kwargs={"k": 3, "filter": filter_dict}
    )
    docs = retriever.invoke(query)
    if not docs:
        return "未找到匹配的题目，请尝试其他关键词"

    result = []
    for doc in docs:
        topic = doc.metadata.get("topic", "未分类")
        page = doc.page_content
        # 把"问题:"和"答案:"拆开，便于 Agent 区分
        result.append(
            f"【题目-{topic}】{page}\n"
            f" 上面的'答案'部分是给你内部评估用的，面试时只问'问题'部分，不要念答案！"
        )
    return "\n\n---\n\n".join(result)

#评估回答工具（仅用于 Agent 内部决策，结果不暴露给用户）
@tool
def evaluate_answer(question: str, user_answer: str, reference_answer: str) -> str:
    """
    内部评估候选人的回答。返回评分和决策建议。
    注意：Agent 拿到结果后自己决定追问/换题/结束，不要把评分透露给候选人！
    """
    llm = get_eval_llm()
    prompt = f"""你是面试官内部评估系统。根据以下信息判断下一步动作。

题目：{question}
参考答案：{reference_answer}
候选人回答：{user_answer}

请严格按此格式输出（三行，不要多写）：
动作：追问/下一题/结束
方向：追问方向或下一题话题建议

判断规则（按优先级）：
1. 候选人明确表示"不知道""不太清楚""没了解过""不会""没接触过""不太懂" → 动作:下一题，方向:换一个完全不同的技术领域
2. 候选人回答深入、有见地 → 随机决定，约40%追问深入细节，约60%换下一题
3. 候选人回答基本正确但浅 → 动作:下一题，方向:换相关但不完全相同的方向
4. 不要连续给"追问"——如果上一轮已经追问过了，这次必须下一题
5. 只有明确要求结束时才用"结束"动作"""
    response = llm.invoke(prompt)
    return response.content

#调整难度工具
@tool
def adjust_difficulty(scores:str) -> str:
    """
    根据历史得分调整难度.输入格式"8,7,9".
    返回"easy","medium"或"hard".
    """
    try:
        score_list = [int(s.strip()) for s in scores.split(",")if s.strip()]
        if not score_list:
            return "medium"
        avg = sum(score_list) / len(score_list)
        if avg >= 8:
            return "hard"
        elif avg >=5:
            return "medium"
        else:
            return "easy"
    except:
        return "medium"

#生成面试报告工具
@tool
def generate_final_report(history: str) -> str:
    """
    根据完整面试历史生成结构化评估报告
    history格式:每轮:"问题:xxx\n回答:xxx\n评分:xxx\n反馈:xxx"
    """
    llm =get_eval_llm()
    prompt = f"""请根据以下面试记录,生成一份结构化评估报告.

面试记录:
{history}

报告格式:
总体评分(1-10)
技术能力分析(强项/弱项)
沟通与表达
录用建议(强烈推荐/推荐/待定/不推荐)
后续建议
"""
    response = llm.invoke(prompt)
    return response.content

#构建agent
def build_agent(resume_text: str, job_desc: str, difficulty: str, max_questions: int):
    """构建带 4 个工具的 Agent，注入面试上下文"""
    tools = [
        search_question_bank,
        evaluate_answer,
        adjust_difficulty,
        generate_final_report,
    ]
    llm = get_llm(temperature=0.7)
    prompt = ChatPromptTemplate.from_messages([
        ("system", f"""你是一位资深技术面试官，正在进行一场真实的技术面试。

【候选人信息】
简历：{resume_text}
岗位：{job_desc}
当前难度：{difficulty}
计划问题数：{max_questions}

【面试规则 —— 严格遵守】
1. 每次只问一道题，不要一次抛出多个问题
2. 候选人回答后，调用 evaluate_answer 做内部评估（你需要提供 question, user_answer, reference_answer 三个参数）
3. evaluate_answer 会返回评分和动作建议，你据此决定下一步：
   - "追问" → 基于候选人刚才的回答，深入追问一个更细节的问题
   - "下一题" → 调用 search_question_bank 换一个不同方向的新题
   - "结束" → 进入面试总结阶段
4.  换题时 search_question_bank 必须用不同的 query！不要用同样的关键词再搜一遍
5. 问问题的时候只问问题,不需要说"这是哪个知识点的问题",也不用说些其他的什么内容,专注问题就好,可适当说一些"嗯" "好"这些语气词,你只能说问题
【不知道就跳过 —— 铁律】
- 候选人说"不知道""不太清楚""没了解过""不会""没接触过"等等表示回答不了这个问题的回答时 → 不要再追问同一个话题！
- 如果是简单题回答不出来那么把这个情况参考到最后的是否录用中(适当扣最后的总分) ,如果是难题就跳过不影响最后的成绩(不需要调用 evaluate_answer)
- 立刻调用 search_question_bank，用完全不同的关键词搜新题（比如之前问并发，现在就搜集合、JVM、设计模式等）
- 不要试图引导、暗示、鼓励——立刻跳过去问别的

【回答正确时 —— 随机应变】
- 回答正确时，随机决定下一步（不要总是一个模式）：
  - 约 40% 概率：追问一个更深的问题（考察深度）
  - 约 60% 概率：直接换下一道题（考察知识广度）
- 避免连续追问超过 2 次

【重要禁忌 —— 违反即面试事故】
- 面试过程中绝对不要把评分透露给候选人！评分是内部参考
- 绝对不要把参考答案展示给候选人！search_question_bank 返回的"答案"部分是给你自己看的，你只需要问"问题"部分
- 不要在问题里包含"参考答案""标准答案"这些字眼
- 不要对候选人的回答做长篇点评——你只需要决定追问还是下一题
- 不要一次问多个问题
- 当前问题结束直接问下一个问题,不要说其他内容
- 不要问简历里没提到的技术（除非是岗位基本要求）

【结束时机】
- 累计问了 {max_questions} 道题（追问不算）后，调用 generate_final_report 生成完整报告
- 报告生成后，把报告内容展示给候选人

用中文交流，语气专业但友好。像真实面试官一样自然对话。"""),
        ("user", "{input}"),
        MessagesPlaceholder(variable_name="agent_scratchpad"),
    ])

    agent = create_tool_calling_agent(llm, tools, prompt)
    return AgentExecutor(agent=agent, tools=tools, verbose=True)


# ============================================================
# Streamlit 界面
# ============================================================
def run_app():
    st.set_page_config(page_title="AI 面试官", layout="wide")
    st.title("AI 面试官")


    #记录每一轮问题 回答 得分
    if "interview_history" not in st.session_state:
        st.session_state.interview_history = [] # 每轮记录 {question, answer, score, feedback}
    if "current_difficulty" not in st.session_state:
        st.session_state.current_difficulty = "medium"
    if "question_count" not in st.session_state:
        st.session_state.question_count = 0
    if "max_question_count" not in st.session_state:
        st.session_state.max_question_count = 10
    if "phase" not in st.session_state:
        st.session_state.phase = "idle"
    if "current_question_ref" not in st.session_state:
        st.session_state.current_question_ref = None
    if "agent" not in st.session_state:
        st.session_state.agent = None
    # 初始化 session_state
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "current_question" not in st.session_state:
        st.session_state.current_question = None

    # ===================== 侧边栏 =====================
    with st.sidebar:
        st.header("求职者信息")
        upload_file = st.file_uploader("上传简历", type=["pdf", "docx", "txt"])
        job_desc = st.text_input("应聘岗位", placeholder="例如：Java 开发工程师")

        if st.button(" 开始面试"):
            if upload_file:
                resume_text = parse_resume(upload_file)
                st.session_state.resume_text = resume_text
                st.session_state.job_desc = job_desc
                st.session_state.messages = []
                st.session_state.interview_history = []
                st.session_state.current_difficulty = "medium"
                st.session_state.question_count = 0
                st.session_state.current_question = None
                st.session_state.phase = "questioning"
                st.session_state.retrieved_docs = None
                # 把简历上下文注入 Agent
                st.session_state.agent = build_agent(
                    resume_text=resume_text,
                    job_desc=job_desc,
                    difficulty="medium",
                    max_questions=st.session_state.max_question_count,
                )
                st.rerun()
            else:
                st.warning("请先上传简历")

        # 自动出题
        if st.session_state.get("phase") == "questioning" and not st.session_state.get("current_question"):
            with st.spinner(" 面试官正在出题..."):
                agent = st.session_state.agent
                progress = f"[面试进度] 已问 0/{st.session_state.max_question_count} 题 | 难度: {st.session_state.current_difficulty}"
                result = agent.invoke({
                    "input": f"{progress}\n请开始面试，先调用 search_question_bank 找一道合适的题来问。"
                })
                st.session_state.messages.append({"role": "assistant", "content": result["output"]})
                st.session_state.current_question = result["output"]
                st.rerun()

    # ===================== 主界面：聊天记录 =====================
    for msg in st.session_state.get("messages", []):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # ===================== 用户输入框 =====================
    if prompt := st.chat_input(" 请输入你的回答..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.session_state.question_count += 1

        with st.chat_message("assistant"):
            with st.spinner(" 面试官正在思考..."):
                agent = st.session_state.agent
                # 拼接最近对话历史，让 Agent 有记忆
                recent_history = "\n".join([
                    f"{'面试官' if m['role'] == 'assistant' else '候选人'}：{m['content']}"
                    for m in st.session_state.messages[-6:]  # 最近 3 轮对话
                ])
                progress = f"[面试进度] 已问 {st.session_state.question_count}/{st.session_state.max_question_count} 题 | 难度: {st.session_state.current_difficulty}"
                result = agent.invoke({
                    "input": f"{progress}\n\n【对话历史】\n{recent_history}\n\n候选人刚才说：{prompt}\n\n请根据你的判断决定下一步：追问 / 下一题 / 总结。"
                })
                st.markdown(result["output"])

        st.session_state.messages.append({"role": "assistant", "content": result["output"]})
        st.session_state.current_question = result["output"]

        # 检测面试是否结束
        if "总体评分" in result["output"] or "评估报告" in result["output"]:
            st.session_state.phase = "done"

        st.rerun()

if __name__ == "__main__":
    run_app()